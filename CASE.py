import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import base64
import json
from openai import OpenAI
import os
from datetime import datetime
import io

apikey = st.secrets["API_KEY"]
client = OpenAI(api_key = apikey)

def prompt(subject, topic, language):
    if language == 'German':
        prompt_text = f"Create a worksheet for the subject {subject} on {topic}. The first 5 questions should be comprehension questions. The second 2 questions should be multiple-choice questions (a), b), c), d)), and the last question should be a cloze (approximately 4 sentences). It should be in this format: worksheet: ['Heading', 'Comprehension question', 'Comprehension question', 'Comprehension question', 'Comprehension question', 'Comprehension question 5', 'Multiple-choice question a) answer, b) answer, c) answer, d) answer', 'Multiple-choice', 'Cloze']"
    elif language == 'English':
        prompt_text = f"Create a worksheet for the subject {subject} on {topic}. The first 5 questions should be comprehension questions. The second 2 questions should be multiple-choice questions (a), b), c), d)), and the last question should be a cloze (approximately 4 sentences). It should be in this format: worksheet: ['Heading', 'Comprehension question', 'Comprehension question', 'Comprehension question', 'Comprehension question', 'Comprehension question 5', 'Multiple-choice question a) answer, b) answer, c) answer, d) answer', 'Multiple-choice', 'Cloze']"
    elif language == 'French':
        prompt_text = f"Create a worksheet for the subject {subject} on {topic}. The first 5 questions should be comprehension questions. The second 2 questions should be multiple-choice questions (a), b), c), d)), and the last question should be a cloze (approximately 4 sentences). It should be in this format: worksheet: ['Heading', 'Comprehension question', 'Comprehension question', 'Comprehension question', 'Comprehension question 5', 'Multiple-choice question a) answer, b) answer, c) answer, d) answer', 'Multiple-choice', 'Cloze']"
    return prompt_text

def response(prompt_text):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo-1106",
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": "You are a helpful assistant designed to output JSON."},
            {"role": "user", "content": prompt_text},
        ]
    )
    while True:
        try:
            worksheet_list = json.loads(response.choices[0].message.content)["worksheet"]
            break
        except:
            pass
    return worksheet_list

def word(worksheet, subject):
    doc = Document()
    title = worksheet[0].replace(" ", "-").lower()

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri Light'
    font.size = Pt(12)
    font.color.rgb = RGBColor(40, 40, 40)
    
    current_date = datetime.now()
    date = current_date.strftime("%d.%m.%y")
    header_style = doc.styles["Header"]
    header_paragraph = doc.sections[0].header.paragraphs[0]
    header_paragraph.text = f"{subject}\t\t{date}"
    header_paragraph.style = header_style

    header_style = doc.styles["Heading1"] 
    header_paragraph = doc.add_paragraph()
    header_run = header_paragraph.add_run(worksheet[0])
    header_run.bold = True
    header_run.font.size = Pt(18) 

    for _ in range(2):
        doc.add_paragraph()

    for i, content in enumerate(worksheet[1:], start=1):
        subtitle_paragraph = doc.add_paragraph(f"{i}) {content}")
        subtitle_run = subtitle_paragraph.runs[0]
        subtitle_run.font.bold = True
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(0, 28, 46)
        for _ in range(3):
            doc.add_paragraph()

    doc_name = f"{title}.docx"
  
    doc.save(doc_name)
    return doc_name

def create_worksheet(subject_selection, topic, subject, language):
    worksheet = response(prompt(subject, topic, language))
    doc_download = word(worksheet, subject_selection)
    return doc_download

st.set_page_config(
    page_title="CASE",
    page_icon="🏫",
)

st.title("CASE")

subject_selection = st.selectbox(
    "Subject",
    ["Select a Subject"] + ["Enter Your Own Subject"] + ["Mathematics 🔢", "German 📚", "English 🇬🇧", "History 📜", "Geography 🌍", "Biology 🌿", "Chemistry 🧪", "Physics ⚙️", "Computer Science 💻", "Music 🎵", "Art 🎨", "Sports 🏃‍♂️", "Ethics 🤔", "Religion ⛪", "Politics 🗳️", "Economics 💹", "Philosophy 🤯", "Social Studies 👥", "Psychology 🧠", "Sociology 👩‍👩‍👧‍👦", "Foreign Language 🗣️", "Latin 🏛️", "Spanish 🇪🇸", "French 🇫🇷", "Italian 🇮🇹", "Russian 🇷🇺", "Chinese 🇨🇳", "Japanese 🇯🇵", "Korean 🇰🇷", "Arabic 🇸🇦", "Media Studies 📱"],
    key="subject_dropdown"
)


if subject_selection != "Select a Subject":
    if subject_selection == "Enter Your Own Subject":
        subject_selection = st.text_input("Enter Your Own Subject")[:-1]
        st.empty()

topic = st.text_input("Topic:")

language_options = ['German🇩🇪', 'English🇬🇧', 'French🇫🇷']

language = st.selectbox('Choose Your Language:', language_options)

create_button = st.button("Create Worksheet")


if topic and create_button:
    doc_download = create_worksheet(subject_selection, topic, subject_selection, language)
    bio = io.BytesIO()
    doc = Document(doc_download)
    doc.save(bio)

    st.success("Worksheet created successfully!")

    st.download_button(
        label="Click here to download",
        data=bio.getvalue(),
        file_name="Worksheet.docx",
        mime="docx",
        key="download_button",
        help="green"
    )

